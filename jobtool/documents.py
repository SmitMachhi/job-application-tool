from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_resume_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if name.endswith(".pdf"):
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    raise ValueError("Unsupported resume format")


def save_docx(markdownish_text: str, path: Path) -> None:
    doc = Document()
    for block in markdownish_text.split("\n"):
        line = block.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
